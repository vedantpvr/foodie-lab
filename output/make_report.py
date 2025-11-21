#!/usr/bin/env python3
"""
make_report.py
Generates a professional DOCX report for the Firebase Recipe Analytics Pipeline
with clear placeholders where images should be inserted manually.

Requirements:
    pip install python-docx

Run:
    python make_report.py
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# Title page
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run("Firebase Recipe Analytics Pipeline — Data Engineering Report\n")
run.bold = True
run.font.size = Pt(20)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
meta.add_run("Author: Sanket Raut\n").bold = True
meta.add_run("Date: November 2025\n").italic = True
meta.add_run("Version: 1.0\n")

doc.add_page_break()

def add_heading(text):
    h = doc.add_heading(level=1)
    h_run = h.add_run(text)
    h_run.font.size = Pt(14)
    return h

# 1. Executive Summary
add_heading("1. Executive Summary")
p = doc.add_paragraph()
p.add_run("This report documents the complete Data Engineering pipeline developed for Firebase-based Recipe Analytics. "
          "It ingests data from Firestore, normalizes it into CSV files, validates data quality, and produces analytics insights and visualizations.")

# 2. Data Model
add_heading("2. Data Model")
doc.add_paragraph("The model has five primary entities: Users, Recipes, Ingredients, Steps, Interactions.")
doc.add_paragraph("PLACEHOLDER: Insert ERD image here")
doc.add_paragraph("[INSERT IMAGE: ERD.png]")

# 3. System Architecture
add_heading("3. System Architecture")
doc.add_paragraph("Extraction → Transformation → Validation → Analytics → Reporting")
doc.add_paragraph("Mermaid representation (for reference):")
doc.add_paragraph("graph LR\nA[Firestore] --> B[Node.js ETL Scripts]\nB --> C[Normalized CSV Files]\nC --> D[Validation Engine]\nD --> E[Analytics]\nE --> F[Reports & Charts]")

# 4. ETL Pipeline
add_heading("4. ETL Pipeline Description")
doc.add_paragraph("4.1 Extraction: Firebase Admin SDK (Node.js)")
doc.add_paragraph("4.2 Transformation: Flatten nested subcollections into relational CSVs")
doc.add_paragraph("4.3 Loading: Write CSVs to output/ folder")
doc.add_paragraph("Output files:")
doc.add_paragraph("- recipe.csv\n- ingredients.csv\n- steps.csv\n- interactions.csv\n- users.csv")

# 5. Data Quality Validation
add_heading("5. Data Quality Validation")
doc.add_paragraph("Rules applied:")
doc.add_paragraph("- Required fields must be present")
doc.add_paragraph("- Numeric fields must be non-negative (or positive where applicable)")
doc.add_paragraph("- Difficulty must be one of: easy, medium, hard")
doc.add_paragraph("- Ratings between 0 and 5")
doc.add_paragraph("- Each recipe should have at least 1 ingredient and 1 step")
doc.add_paragraph("Validation example summary:")
tbl = doc.add_table(rows=1, cols=4)
hdr = tbl.rows[0].cells
hdr[0].text = "Table"; hdr[1].text = "Total Rows"; hdr[2].text = "Valid"; hdr[3].text = "Invalid"
rows = [
    ("recipes","20","20","0"),
    ("ingredients","200","198","2"),
    ("steps","150","150","0"),
    ("interactions","400","400","0"),
    ("users","20","20","0"),
]
for r in rows:
    cells = tbl.add_row().cells
    cells[0].text = r[0]; cells[1].text = r[1]; cells[2].text = r[2]; cells[3].text = r[3]

doc.add_page_break()

# 6. Recipe Analytics
add_heading("6. Recipe Analytics (Charts)")
doc.add_paragraph("6.1 Preparation Time Distribution")
doc.add_paragraph("[INSERT IMAGE: prep_time_histogram.png]")
doc.add_paragraph("6.2 Prep Time vs Like Count")
doc.add_paragraph("[INSERT IMAGE: prep_vs_likes_scatter.png]")
doc.add_paragraph("6.3 Top Ingredients (by frequency)")
doc.add_paragraph("[INSERT IMAGE: top_ingredients.png]")

# 7. User Analytics
add_heading("7. User Analytics (Charts)")
doc.add_paragraph("7.1 Users by Country")
doc.add_paragraph("[INSERT IMAGE: users_by_country.png]")
doc.add_paragraph("7.2 Top Users by Interaction Count")
doc.add_paragraph("[INSERT IMAGE: top_users_by_interactions.png]")

# 8. Insights Summary
add_heading("8. Insights Summary")
doc.add_paragraph("- Most common ingredients include ghee, turmeric, wheat flour, and jaggery.")
doc.add_paragraph("- Prep time shows a weak-to-moderate correlation with likes.")
doc.add_paragraph("- A small set of users drive most engagement.")
doc.add_paragraph("- Majority of users are located in India (IN).")

# 9. Output Summary
add_heading("9. Output Summary")
doc.add_paragraph("Files produced in output/: recipe.csv, ingredients.csv, steps.csv, interactions.csv, users.csv, validation_report.csv, analytics_summary.txt, charts/*")

# 10. Conclusion
add_heading("10. Conclusion & Next Steps")
doc.add_paragraph("The ETL pipeline is complete and produces validated CSV artifacts. Next steps: dashboard, BigQuery integration, recommendation models.")

doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
f.add_run("Prepared for team review — Sanket Raut").italic = True

# Save
doc.save("Firebase_Recipe_Analytics_Report.docx")
print("Saved: Firebase_Recipe_Analytics_Report.docx")
